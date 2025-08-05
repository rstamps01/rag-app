"""
Enhanced Documents API Routes
Provides document management with vector database integration
"""

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from sqlalchemy.orm import Session
from typing import List, Optional
import os
import uuid
import time
import logging
from datetime import datetime

from app.db.session import get_db
from app.models.models import Document
from app.services.enhanced_vector_db import enhanced_vector_db_service
from app.core.config import settings

logger = logging.getLogger(__name__)

router = APIRouter()

# Document processing utilities
def extract_text_from_file(file_path: str, content_type: str) -> str:
    """Extract text content from uploaded file"""
    try:
        if content_type == "text/plain":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif content_type == "application/pdf":
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                logger.warning("PyPDF2 not available, using basic text extraction")
                return f"PDF file: {os.path.basename(file_path)}"
        
        elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                logger.warning("python-docx not available, using basic text extraction")
                return f"Word document: {os.path.basename(file_path)}"
        
        else:
            # For other file types, return filename as content
            return f"Document: {os.path.basename(file_path)}"
            
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {e}")
        return f"Document: {os.path.basename(file_path)}"

@router.get("/")
async def get_documents(
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=1000),
    department: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    db: Session = Depends(get_db)
):
    """Get documents with filtering and pagination"""
    try:
        logger.info(f"Documents requested: skip={skip}, limit={limit}, department={department}, status={status}")
        
        # Build query
        query = db.query(Document)
        
        # Apply filters
        if department:
            query = query.filter(Document.department == department)
        if status:
            query = query.filter(Document.status == status)
        
        # Get total count
        total = query.count()
        
        # Apply pagination
        documents = query.offset(skip).limit(limit).all()
        
        # Format response
        doc_list = []
        for doc in documents:
            doc_list.append({
                "id": doc.id,
                "filename": doc.filename,
                "upload_date": doc.upload_date.isoformat() if doc.upload_date else "",
                "size": doc.size or 0,
                "status": doc.status,
                "department": doc.department,
                "content_type": doc.content_type,
                "processing_status": doc.processing_status,
                "vector_stored": doc.vector_stored
            })
        
        return {
            "documents": doc_list,
            "total": total,
            "skip": skip,
            "limit": limit,
            "filters": {
                "department": department,
                "status": status
            },
            "source": "database"
        }
        
    except Exception as e:
        logger.error(f"Failed to get documents: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve documents: {str(e)}")

@router.post("/")
async def upload_document(
    file: UploadFile = File(...),
    department: Optional[str] = Form("General"),
    process_immediately: bool = Form(True),
    db: Session = Depends(get_db)
):
    """Upload a document with optional immediate processing"""
    try:
        logger.info(f"Document upload: {file.filename}, department: {department}")
        
        # Validate file type
        allowed_extensions = {".pdf", ".txt", ".docx", ".md", ".doc"}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_ext} not allowed. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Read file content
        content = await file.read()
        file_size = len(content)
        
        # Check file size (100MB limit)
        max_size = 100 * 1024 * 1024
        if file_size > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"File size ({file_size // (1024*1024)}MB) exceeds limit (100MB)"
            )
        
        # Generate unique filename and ID
        file_id = str(uuid.uuid4())
        unique_filename = f"{file_id}{file_ext}"
        
        # Create upload directory if it doesn't exist
        upload_dir = "/app/data/uploads"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, unique_filename)
        
        # Save file
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Create document record
        document = Document(
            id=file_id,
            filename=file.filename,
            content_type=file.content_type,
            size=file_size,
            status="uploaded",
            path=file_path,
            department=department,
            processing_status="pending" if process_immediately else "queued",
            vector_stored=False
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document saved: {file_path}, ID: {file_id}")
        
        # Process document immediately if requested
        processing_result = None
        if process_immediately:
            processing_result = await process_document_for_vector_db(file_id, db)
        
        return {
            "id": file_id,
            "filename": file.filename,
            "size": file_size,
            "status": "uploaded",
            "department": department,
            "upload_time": time.time(),
            "processing_requested": process_immediately,
            "processing_result": processing_result,
            "message": "Document uploaded successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.post("/{document_id}/process")
async def process_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Process a document for vector database storage"""
    return await process_document_for_vector_db(document_id, db)

async def process_document_for_vector_db(document_id: str, db: Session) -> dict:
    """Process document for vector database storage"""
    try:
        # Get document from database
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        logger.info(f"Processing document {document_id} for vector storage")
        
        # Update processing status
        document.processing_status = "processing"
        db.commit()
        
        # Extract text content
        text_content = extract_text_from_file(document.path, document.content_type)
        
        if not text_content or len(text_content.strip()) < 10:
            document.processing_status = "failed"
            document.status = "failed"
            db.commit()
            return {
                "success": False,
                "message": "Failed to extract meaningful text content",
                "document_id": document_id
            }
        
        # Process with vector database
        if enhanced_vector_db_service.is_available():
            success = enhanced_vector_db_service.process_document(
                document_id=document_id,
                filename=document.filename,
                content=text_content,
                department=document.department
            )
            
            if success:
                document.processing_status = "completed"
                document.status = "processed"
                document.vector_stored = True
                logger.info(f"✅ Document {document_id} processed successfully")
            else:
                document.processing_status = "failed"
                document.status = "failed"
                logger.error(f"❌ Vector processing failed for document {document_id}")
        else:
            document.processing_status = "failed"
            document.status = "failed"
            logger.error("Vector database service not available")
        
        db.commit()
        
        return {
            "success": document.vector_stored,
            "message": "Document processed successfully" if document.vector_stored else "Vector processing failed",
            "document_id": document_id,
            "status": document.status,
            "processing_status": document.processing_status,
            "vector_stored": document.vector_stored,
            "text_length": len(text_content)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document processing error: {e}")
        
        # Update document status on error
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if document:
                document.processing_status = "failed"
                document.status = "failed"
                db.commit()
        except:
            pass
        
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")

@router.get("/{document_id}")
async def get_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Get document details"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "id": document.id,
            "filename": document.filename,
            "upload_date": document.upload_date.isoformat() if document.upload_date else "",
            "size": document.size,
            "status": document.status,
            "department": document.department,
            "content_type": document.content_type,
            "processing_status": document.processing_status,
            "vector_stored": document.vector_stored,
            "path": document.path
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve document: {str(e)}")

@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    db: Session = Depends(get_db)
):
    """Delete a document and its vector data"""
    try:
        # Get document from database
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete from vector database if stored
        if document.vector_stored and enhanced_vector_db_service.is_available():
            enhanced_vector_db_service.delete_document(document_id)
        
        # Delete file from filesystem
        if document.path and os.path.exists(document.path):
            os.remove(document.path)
            logger.info(f"File deleted: {document.path}")
        
        # Delete from database
        db.delete(document)
        db.commit()
        
        logger.info(f"Document {document_id} deleted successfully")
        
        return {
            "success": True,
            "message": "Document deleted successfully",
            "document_id": document_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")

@router.get("/stats/overview")
async def get_documents_stats(db: Session = Depends(get_db)):
    """Get document statistics"""
    try:
        total_documents = db.query(Document).count()
        processed_documents = db.query(Document).filter(Document.status == "processed").count()
        failed_documents = db.query(Document).filter(Document.status == "failed").count()
        pending_documents = db.query(Document).filter(Document.status == "uploaded").count()
        
        # Get department breakdown
        department_stats = db.query(
            Document.department,
            db.func.count(Document.id).label('count')
        ).group_by(Document.department).all()
        
        # Get vector storage stats
        vector_stored = db.query(Document).filter(Document.vector_stored == True).count()
        
        return {
            "total_documents": total_documents,
            "processed_documents": processed_documents,
            "failed_documents": failed_documents,
            "pending_documents": pending_documents,
            "vector_stored": vector_stored,
            "department_breakdown": [
                {"department": dept, "count": count} 
                for dept, count in department_stats
            ],
            "vector_db_info": enhanced_vector_db_service.get_collection_info()
        }
        
    except Exception as e:
        logger.error(f"Failed to get document stats: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve statistics: {str(e)}")

